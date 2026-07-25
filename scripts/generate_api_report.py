"""
Genera el informe "API Inteligente: endpoints claros y testeables"
del AI Phishing Detector como .docx
Salida: docs/informe_api_endpoints.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("API Inteligente")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Exponer el proyecto mediante endpoints claros y testeables")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. INVENTARIO DE ENDPOINTS ──────────────────────────────────────────
heading(doc, "1. Inventario de Endpoints", 1, "1E3A5F")
p = doc.add_paragraph(
    "El backend expone 6 endpoints REST sobre FastAPI, agrupados en /docs (Swagger) bajo "
    "3 tags: Sistema, Analisis y Cache. Todos menos / y /health estan detras de "
    "require_api_key (header X-API-Key); /predict y /analyze-content ademas estan sujetos "
    "a rate limiting (30 solicitudes / 60s por IP)."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Metodo", "Ruta", "Tag", "Auth", "Rate limit", "Descripcion"],
    [
        ["GET",    "/",               "Sistema",  "No",  "No",  "Sanity check de la raiz"],
        ["GET",    "/health",         "Sistema",  "No",  "No",  "Liveness check (usado por la extension y el HEALTHCHECK de Docker)"],
        ["POST",   "/predict",        "Analisis", "Si",  "Si",  "Pipeline completo de analisis de URL (WHOIS, HTML, 3 APIs, 3 modelos ML)"],
        ["POST",   "/analyze-content","Analisis", "Si",  "Si",  "Clasificacion de texto libre REAL/FAKE"],
        ["GET",    "/cache/stats",    "Cache",    "Si",  "No",  "Estadisticas del cache en memoria (entries, valid, ttl, max_size)"],
        ["DELETE", "/cache",          "Cache",    "Si",  "No",  "Evictar todos los resultados cacheados"],
    ],
    col_widths=[1.4, 2.5, 1.4, 1.1, 1.4, 6.2],
)

# ── 2. CONTRATOS DE REQUEST/RESPONSE ─────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Contratos de Entrada y Salida", 1, "1E3A5F")

heading(doc, "2.1 POST /predict", 2, "1565C0")
bullet(doc, "Request: UrlRequest { url: str } — Pydantic valida que empiece con http:// o https://; si no, 422 automatico")
bullet(doc, "response_model=PredictResponse (backend/app/schemas/response_schema.py): tipa el top-level; los sub-diccionarios (risk_assessment, machine_learning, html_analysis, etc.) quedan como dict abierto a proposito, porque cada senal puede degradarse a {\"error\": \"...\"} si su sub-servicio falla — tiparlos de forma estricta rompería la respuesta ante cualquier fallo parcial de un proveedor externo")
add_code_block(doc, '// Request\n{ "url": "https://ejemplo.com" }')
add_code_block(doc,
"""// Response 200 (forma real, capturada en tests/test_api.py)
{
  "url": "https://ejemplo.com",
  "risk_assessment": { "risk": "LOW", "confidence": 90, "score": 90, "reasons": [...] },
  "machine_learning": {
    "fusion":        { "prediction": 0, "phishing_probability": 0.12, "legitimate_probability": 0.88, "rf_weight": 0.4, "roberta_weight": 0.6 },
    "random_forest": { "prediction": 0, "phishing_probability": 0.13, "legitimate_probability": 0.87 },
    "roberta":       { "prediction": 1, "phishing_probability": 0.11, "legitimate_probability": 0.89 }
  },
  "html_analysis": { "success": true, "html_features": {...} },
  "url_features":  { "has_https": true, "url_length": 22, "num_dots": 2, "has_ip": false },
  "domain_info":   { "domain": "ejemplo.com", "tld": "com", "domain_age_days": 10000 },
  "cached": false,
  "analysis_time_ms": 1847,
  "virustotal": { "verdict": "clean" },
  "safe_browsing": { "is_threat": false },
  "fact_check": { "verdict": "reliable" },
  "content_classification": null
}""")
bullet(doc, "Errores: 422 si falta url o el body es invalido; 403 si falta/es invalido X-API-Key (con API_KEY configurada); 429 si se excede el rate limit")

doc.add_paragraph()
heading(doc, "2.2 POST /analyze-content", 2, "1565C0")
bullet(doc, "Request: TextRequest { text: str } — el campo ahora documenta el umbral de 300 caracteres via Field(description=...); la validacion en si sigue viviendo en ContentClassifierService, no en el schema (es una decision de negocio, no de formato, para no romper el flujo actual)")
bullet(doc, "response_model=ContentAnalysisResponse: todos los campos son opcionales porque el endpoint tiene 3 formas de respuesta validas (clasificado / no_content / error)")
add_code_block(doc, '// Request\n{ "text": "..." }  // >= 300 caracteres para ser clasificado')
add_code_block(doc,
"""// Response 200 — texto largo
{ "verdict": "real", "label": "REAL", "confidence": 0.93, "raw_label": "REAL" }

// Response 200 — texto corto (< 300 chars), no se clasifica
{ "verdict": "no_content", "label": "UNKNOWN", "confidence": 0.0 }""")

doc.add_paragraph()
heading(doc, "2.3 GET /health, GET /, GET /cache/stats, DELETE /cache", 2, "1565C0")
add_code_block(doc,
"""GET /health         -> { "status": "healthy" }
GET /               -> { "message": "AI Phishing Detector Running" }
GET /cache/stats     -> { "entries": N, "valid": N, "ttl_seconds": 600, "max_size": 500 }
DELETE /cache        -> { "cleared": N }""")

# ── 3. SEGURIDAD Y CONTROL DE ACCESO ────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Seguridad y Control de Acceso", 1, "1E3A5F")
add_table(doc,
    ["Mecanismo", "Implementacion"],
    [
        ["Autenticacion", "Header X-API-Key, verificado por require_api_key (dependency de FastAPI aplicada a todo el router). Si API_KEY esta vacio en .env, la verificacion es no-op (documentado, no accidental)"],
        ["Rate limiting", "RateLimitMiddleware: 30 solicitudes / 60s por IP, solo sobre /predict y /analyze-content; responde 429 con header Retry-After: 60"],
        ["CORS", "allow_origin_regex acepta chrome-extension://* (32 chars), localhost/127.0.0.1, y origenes extra via ALLOWED_ORIGINS"],
        ["Logging de requests", "RequestLoggingMiddleware registra metodo, path, status, latencia e IP de cada solicitud (incluye las bloqueadas por rate limit)"],
        ["Fail-fast en produccion", "Si ENVIRONMENT=production y API_KEY esta vacio, el backend rehusa arrancar"],
    ],
    col_widths=[3.5, 10.5],
)

# ── 4. TESTEABILIDAD ─────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "4. Testeabilidad", 1, "1E3A5F")
p = doc.add_paragraph(
    "Los endpoints se prueban con FastAPI TestClient, sin depender de red ni modelos reales "
    "(model loaders y servicios externos se mockean via conftest.py y unittest.mock.patch)."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Endpoint", "Archivo de tests", "Casos cubiertos", "# tests"],
    [
        ["GET /health, GET /",
         "tests/test_api.py",
         "Status 200, forma de la respuesta",
         "4"],
        ["POST /predict",
         "tests/test_api.py",
         "Validacion 422 (sin url, body invalido, sin body); 200 con mock de PhishingService; estructura completa de risk_assessment/machine_learning/virustotal/safe_browsing/fact_check; escenario seguro vs phishing; eco de la url",
         "12"],
        ["POST /analyze-content",
         "tests/test_content_api.py",
         "Validacion 422 (sin text, campo incorrecto, sin body); clasificacion REAL/FAKE mockeada; umbral de 300 caracteres (limites 50/299/300); propagacion de errores del servicio",
         "16"],
        ["GET /cache/stats, DELETE /cache",
         "tests/test_cache_and_security.py",
         "Status 200 y forma de la respuesta para ambos endpoints",
         "4"],
        ["require_api_key (403)",
         "tests/test_cache_and_security.py",
         "403 sin header X-API-Key y con key invalida cuando API_KEY esta configurada; 200 con la key correcta",
         "3"],
        ["RateLimitMiddleware (429)",
         "tests/test_cache_and_security.py",
         "200 hasta el limite (30), 429 + header Retry-After al superarlo",
         "1"],
    ],
    col_widths=[3.3, 3.3, 5.9, 1.5],
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Total suite del proyecto: ")
run.bold = True
run.font.size = Pt(9.5)
run2 = p.add_run("265 tests, 265 aprobados (pytest.ini: testpaths = tests backend/tests).")
run2.font.size = Pt(9.5)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Nota de diseno de tests: ")
run.bold = True
run.font.size = Pt(9.5)
run2 = p.add_run(
    "RateLimitMiddleware cuenta solicitudes por IP, no por endpoint — un fixture "
    "autouse en tests/conftest.py resetea ese contador global antes de cada test "
    "para que la suite no dependa del orden de ejecucion de los archivos."
)
run2.font.size = Pt(9.5)

# ── 5. ESTADO DE LAS RECOMENDACIONES ─────────────────────────────────────
doc.add_page_break()
heading(doc, "5. Estado de las Recomendaciones", 1, "1E3A5F")
p = doc.add_paragraph(
    "Las brechas identificadas en la version anterior de este informe fueron "
    "implementadas y verificadas contra la suite de tests."
)
p.runs[0].font.size = Pt(9.5)

doc.add_paragraph()
heading(doc, "5.1 Claridad de la API (OpenAPI / Swagger) — implementado", 2, "2E7D32")
bullet(doc, "response_model en los 6 endpoints (backend/app/schemas/response_schema.py: PredictResponse, ContentAnalysisResponse, CacheStatsResponse, CacheClearResponse, HealthResponse, RootResponse)")
bullet(doc, "tags=[\"Sistema\" | \"Analisis\" | \"Cache\"] en cada ruta, mas descripcion de cada tag via openapi_tags en el FastAPI(...) de main.py")
bullet(doc, "summary y description en los 6 endpoints, visibles directamente en /docs")
bullet(doc, "TextRequest.text y UrlRequest.url ahora tienen Field(description=...) documentando el contrato (incluido el umbral de 300 caracteres, que sigue aplicandose en el servicio, no en el schema, para no romper el comportamiento actual)")

doc.add_paragraph()
heading(doc, "5.2 Testeabilidad — implementado", 2, "2E7D32")
bullet(doc, "Nuevo archivo tests/test_cache_and_security.py: GET /cache/stats, DELETE /cache, 403 por API key ausente/invalida, 200 con key correcta, 429 por rate limit excedido")
bullet(doc, "Efecto secundario detectado y corregido: al anadir response_model=PredictResponse, los fixtures mock_safe_response/mock_phishing_response de tests/test_api.py resultaron incompletos frente a la forma real de PhishingService.analyze() (les faltaban cached, analysis_time_ms, virustotal, safe_browsing, fact_check, content_classification) — se completaron para reflejar el contrato real")
bullet(doc, "Efecto secundario detectado y corregido: RateLimitMiddleware comparte un contador global por IP entre /predict y /analyze-content; sumar el nuevo test de 429 exponia una condicion de carrera entre archivos de test (fallaba segun el orden de ejecucion) — se agrego un fixture autouse que resetea el contador antes de cada test")

doc.add_paragraph()
heading(doc, "5.3 Decisiones de diseno no aplicadas literalmente", 2, "1565C0")
bullet(doc, "No se tipo cada sub-diccionario de /predict (risk_assessment, machine_learning, etc.) con sub-modelos Pydantic estrictos: cualquiera de esas senales puede degradarse a {\"error\": \"...\"} si su servicio falla (ver _safe() en phishing_service.py), y un modelo estricto habria provocado 500 (ResponseValidationError) exactamente en el escenario de degradacion elegante que el sistema esta disenado para tolerar")
bullet(doc, "MLPredictionResponse (backend/app/schemas/ml_response.py) se dejo intacto: sigue sin usarse en ningun endpoint por la misma razon; no se elimino porque no fue parte del alcance solicitado")

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — API Inteligente: Endpoints Claros y Testeables  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "informe_api_endpoints.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
