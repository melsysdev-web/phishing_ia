"""
Genera el informe de diagnostico y limpieza (arquitectura, riesgos, plan de accion)
del AI Phishing Detector como .docx
Salida: docs/diagnostico_limpieza_ai_phishing_detector.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    if bold_prefix:
        r1 = p.add_run(f"{bold_prefix}: ")
        r1.bold = True
        r1.font.size = Pt(9.5)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Diagnostico de Arquitectura y Plan de Limpieza")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Componentes  ·  Riesgos  ·  Codigo y archivos sin uso  ·  Plan de accion")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. MAPA DE COMPONENTES ──────────────────────────────────────────────
heading(doc, "1. Mapa de Componentes", 1, "1E3A5F")
p = doc.add_paragraph(
    "Estado verificado contra el codigo actual del repositorio (no contra documentacion "
    "desactualizada). La suite completa de tests pasa: 257/257."
)
p.runs[0].italic = True
p.runs[0].font.size = Pt(9.5)

doc.add_paragraph()
add_code_block(doc,
"""backend/app/          -> pipeline FastAPI (routes -> services -> ml/analyzers/utils)
extension/             -> Manifest V3: popup + sidebar (comparten services/api_client.js)
phishing_detector/     -> modulo standalone, NO conectado al backend (0 imports desde backend/)
training/, scripts/    -> entrenamiento RF y utilidades CLI
docs/                  -> arquitectura, decisiones, entregables semanales (.md + .docx)
tests/, backend/tests/ -> dos suites de pytest, ambas activas via pytest.ini""")

# ── 2. RIESGOS — ALTA CONFIANZA ─────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Riesgos Encontrados", 1, "1E3A5F")

heading(doc, "2.1 Alta confianza — seguro eliminar", 2, "2E7D32")
p = doc.add_paragraph(
    "Verificado con busqueda de referencias en todo el repositorio y analisis estatico "
    "(ruff F401/F811/F841). Nada del codigo activo los usa."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Item", "Motivo"],
    [
        ["estructura.txt / estructura_completa.txt",
         "Volcado de 'tree' en UTF-16, ilegible, se auto-lista a si mismo. Ya excluido en .dockerignore, sin referencias en codigo."],
        ["backend/app/analyzers/test_fetcher.py, test_html_analyzer.py",
         "Scripts de debug ad-hoc (print() + llamada real a google.com), fuera de testpaths de pytest, no documentados."],
        ["Import ContentClassifierService en phishing_service.py:13",
         "Sin uso: /predict nunca clasifica contenido; eso vive solo en /analyze-content."],
        ["Variable page_text en phishing_service.py:75-78",
         "Asignada y nunca leida."],
        ["Variable parsed en feature_mapper.py:10",
         "urlparse(url) calculado y descartado."],
        ["Imports sueltos en generate_architecture_diagram.py y generate_word_report.py",
         "FancyArrowPatch y WD_ALIGN_VERTICAL importados sin usar."],
    ],
    col_widths=[6.0, 8.5],
)

doc.add_paragraph()
heading(doc, "2.2 Requiere decision — no es claramente 'no usado'", 2, "B71C1C")
p = doc.add_paragraph(
    "Son casos de diseno o entregables, no codigo muerto tecnico. Eliminarlos sin confirmar "
    "podria borrar trabajo intencional."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Item", "Contexto"],
    [
        ["phishing_detector/ (modulo completo)",
         "Standalone, documentado como 'not yet wired'. Definir si sigue en desarrollo o se archiva."],
        ["tests/test_{risk_engine,url_features,html_features,feature_mapper}.py vs backend/tests/*",
         "No son duplicados exactos: dos suites distintas para el mismo modulo (fixtures vs constantes inline), ambas aportan cobertura real y ambas corren. Es redundancia de mantenimiento, no codigo muerto."],
        ["docs/*.docx + scripts/generate_*.py que los generan",
         "Parecen entregables de curso (semana 1, arquitectura, seguridad). No tocar sin confirmar que ya fueron entregados."],
        ["extension/content/content.js",
         "No registrado en manifest.json (sin content_scripts) -> hoy no lo carga el navegador. Documentado como placeholder intencional para banner in-page futuro."],
        ["datasets/phishing_detector_sample.csv",
         "Sin referencias en codigo; CSV de ejemplo para correr run_training.py manualmente."],
    ],
    col_widths=[6.0, 8.5],
)

doc.add_paragraph()
heading(doc, "2.3 Sin hallazgos", 2, "1565C0")
bullet(doc, "requirements.txt / backend/requirements.txt estan limpios (verificado import por import contra el codigo real)")
bullet(doc, "extension/services/api_client.js: sus 3 metodos (analyze, analyzeContent, testConnection) se usan en popup.js y sidebar.js")

# ── 3. PLAN DE ACCION ───────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Plan de Accion Propuesto", 1, "1E3A5F")

heading(doc, "Paso 1 — Ejecutar ya (bajo riesgo)", 2, "2E7D32")
bullet(doc, "Borrar estructura.txt, estructura_completa.txt")
bullet(doc, "Borrar backend/app/analyzers/test_fetcher.py y test_html_analyzer.py")
bullet(doc, "Quitar import ContentClassifierService no usado en phishing_service.py")
bullet(doc, "Quitar variable page_text no usada en phishing_service.py")
bullet(doc, "Quitar variable parsed no usada en feature_mapper.py")
bullet(doc, "Quitar imports sueltos en los scripts generate_architecture_diagram.py y generate_word_report.py")

doc.add_paragraph()
heading(doc, "Paso 2 — Pendiente de decision del usuario", 2, "B71C1C")
bullet(doc, "phishing_detector/: mantener como modulo experimental, integrar al pipeline, o eliminar")
bullet(doc, "Consolidar o mantener las dos suites de tests de risk_engine/url_features/html_features/feature_mapper")
bullet(doc, "Confirmar si los .docx en docs/ y sus scripts generadores siguen siendo necesarios")
bullet(doc, "Decidir si content.js se elimina o se deja como placeholder documentado")
bullet(doc, "Confirmar si datasets/phishing_detector_sample.csv se conserva como dataset de ejemplo")

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Diagnostico de Arquitectura y Plan de Limpieza  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "diagnostico_limpieza_ai_phishing_detector.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
