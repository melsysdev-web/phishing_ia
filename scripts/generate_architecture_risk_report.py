"""
Genera el informe de Diagnostico de Arquitectura: Componentes, Riesgos y Plan de
Mejora del AI Phishing Detector como .docx
Salida: docs/diagnostico_arquitectura_riesgos_plan_mejora.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def sev_badge(text, color):
    return (text, color)


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Diagnostico de Arquitectura")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Componentes  ·  Riesgos  ·  Plan de Mejora")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. COMPONENTES ──────────────────────────────────────────────────────
heading(doc, "1. Ubicacion de Componentes", 1, "1E3A5F")
p = doc.add_paragraph(
    "Estado verificado contra el codigo actual del repositorio (rama phishing-ia/risk-engine-v2). "
    "Suite de tests completa: 257/257 aprobados."
)
p.runs[0].italic = True
p.runs[0].font.size = Pt(9.5)

doc.add_paragraph()
heading(doc, "1.1 Capas del sistema", 2, "1565C0")
add_table(doc,
    ["Capa", "Componente", "Ubicacion"],
    [
        ["UI", "Extension Chrome (Popup / Sidebar / Options)", "extension/"],
        ["API", "FastAPI + Uvicorn, CORS, Rate limiting", "backend/app/main.py"],
        ["Seguridad", "X-API-Key header, dependency-level", "backend/app/core/security.py"],
        ["Config", ".env -> settings singleton", "backend/app/core/config.py"],
        ["Rutas", "2 endpoints de analisis + cache ops", "backend/app/api/routes.py"],
        ["Orquestacion", "PhishingService.analyze() + _safe() wrapping", "backend/app/services/phishing_service.py"],
        ["Scoring", "RiskEngine, reglas 0-100 + señales ML", "backend/app/services/risk_engine.py"],
        ["Fusion ML", "FusionEngine, RF 40% + RoBERTa 60%", "backend/app/ml/fusion/fusion_engine.py"],
        ["Modelos", "RandomForest, RoBERTa URL, RoBERTa contenido", "backend/app/random_forest/, roberta/"],
        ["APIs externas", "VirusTotal, Safe Browsing, Fact Check", "backend/app/services/*_service.py"],
        ["Cache", "In-memory TTL 10 min / 500 entradas", "backend/app/utils/url_cache.py"],
        ["Despliegue", "Docker (build-only backend), compose local", "backend/Dockerfile, docker-compose.yml"],
        ["Experimental", "PhishingClassifier fine-tuning standalone", "phishing_detector/ (no conectado)"],
    ],
    col_widths=[2.3, 6.5, 4.8],
)

doc.add_paragraph()
heading(doc, "1.2 Flujo de datos — POST /predict", 2, "1565C0")
add_code_block(doc,
"""Request -> RateLimitMiddleware (30 req/60s por IP) -> require_api_key (X-API-Key)
  |
  v
Cache hit? --si--> respuesta cacheada
  | no
  v
extract_url_features() [CPU]
  |
  v  ThreadPoolExecutor(6 workers)
  WHOIS | HtmlAnalyzer | VirusTotal | SafeBrowsing | FactCheck | RoBERTa URL
  |
  v
FeatureMapper -> RandomForestPredictor (depende de HtmlAnalyzer)
  |
  v
FusionEngine.combine(RF, RoBERTa)  ->  RiskEngine.calculate(...)
  |
  v
Cache.set()  ->  JSON Response""")

# ── 2. RIESGOS ───────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Riesgos Identificados", 1, "1E3A5F")
p = doc.add_paragraph(
    "Riesgos verificados leyendo el codigo real de seguridad, despliegue y operacion "
    "(no documentacion desactualizada). Clasificados por severidad."
)
p.runs[0].font.size = Pt(9.5)

doc.add_paragraph()
heading(doc, "2.1 Seguridad", 2, "B71C1C")
add_table(doc,
    ["Severidad", "Riesgo", "Evidencia"],
    [
        ["ALTA", "Autenticacion desactivada por defecto",
         "require_api_key() es no-op si API_KEY esta vacio en .env (security.py:10-11). Un despliegue sin configurar esa variable queda abierto."],
        ["MEDIA", "Rate limiting no distribuido",
         "_rate_store es un dict en memoria de un solo proceso (main.py:19). Con multiples workers/replicas cada instancia lleva su propio contador; el limite real de 30 req/min se multiplica por N instancias."],
        ["MEDIA", "CORS acepta cualquier extension Chrome",
         "allow_origin_regex admite chrome-extension://[a-z]{32} sin filtrar por ID especifico (main.py:52). Mitigado por X-API-Key, pero cualquier extension que conozca la URL del backend puede intentar llamarlo."],
        ["BAJA", "Sin HTTPS/TLS en el servicio",
         "Uvicorn expone HTTP plano (Dockerfile, docker-compose.yml); en un despliegue publico requeriria un proxy TLS (nginx/Caddy) delante, no incluido."],
        ["BAJA", "Sin logging de auditoria",
         "No hay registro de requests, IPs, ni intentos de auth fallidos mas alla de la respuesta HTTP; dificulta detectar abuso o incidentes."],
    ],
    col_widths=[1.8, 4.7, 7.1],
)

doc.add_paragraph()
heading(doc, "2.2 Escalabilidad y Operacion", 2, "E65100")
add_table(doc,
    ["Severidad", "Riesgo", "Evidencia"],
    [
        ["MEDIA", "Cache volatil en memoria",
         "url_cache.py usa un dict de proceso; se pierde al reiniciar y no se comparte entre replicas (sin Redis ni almacenamiento externo)."],
        ["MEDIA", "Un solo proceso Uvicorn",
         "CMD del Dockerfile no usa --workers ni gunicorn; bajo carga concurrente todo el trafico pasa por un solo proceso Python (el ThreadPoolExecutor(6) ayuda con I/O pero no con CPU-bound como RoBERTa)."],
        ["BAJA", "Sin healthcheck en Docker/Compose",
         "docker-compose.yml no define healthcheck; un contenedor colgado (p.ej. modelo corrupto) no se detecta automaticamente."],
        ["BAJA", "Dependencia de APIs externas con cuota limitada",
         "VirusTotal (tier gratuito ~4 req/min), Safe Browsing y Fact Check tienen cuotas diarias; no hay circuit breaker, solo timeout+catch por servicio (_safe())."],
        ["BAJA", "WHOIS puede ser lento/no confiable",
         "python-whois depende de servidores WHOIS de terceros sin SLA; corre dentro del mismo ThreadPoolExecutor de 6 workers, pudiendo demorar el resto del batch."],
    ],
    col_widths=[1.8, 4.7, 7.1],
)

doc.add_paragraph()
heading(doc, "2.3 Deuda Tecnica y Calidad", 2, "1565C0")
add_table(doc,
    ["Severidad", "Riesgo", "Evidencia"],
    [
        ["BAJA", "Modulo phishing_detector/ desconectado",
         "Fine-tuning standalone completo (model.py, train.py, predict.py) pero PhishingService no lo importa; mantenimiento duplicado si diverge de roberta/."],
        ["BAJA", "Cobertura de tests duplicada",
         "tests/test_{risk_engine,url_features,html_features,feature_mapper}.py y sus pares en backend/tests/ son suites distintas para el mismo modulo; mayor costo de mantenimiento al cambiar una regla de negocio (hay que actualizar dos archivos)."],
        ["BAJA", "Sin CI/CD",
         "No hay workflow de GitHub Actions ni equivalente; tests y lint se ejecutan solo manualmente."],
        ["BAJA", "Sin monitoreo de drift de modelos",
         "No hay mecanismo para detectar cuando la distribucion de URLs de entrada cambia respecto al dataset de entrenamiento."],
    ],
    col_widths=[1.8, 4.7, 7.1],
)

# ── 3. PLAN DE MEJORA ───────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Plan de Mejora", 1, "1E3A5F")

heading(doc, "3.1 Corto plazo (1-2 semanas, bajo esfuerzo)", 2, "2E7D32")
bullet(doc, "Forzar API_KEY obligatoria en despliegues no locales (validar al arrancar que settings.api_key no este vacio si ENV=production)")
bullet(doc, "Agregar HEALTHCHECK a docker-compose.yml apuntando a GET /health")
bullet(doc, "Agregar logging basico de requests (IP, path, status, tiempo) para auditoria minima")
bullet(doc, "Documentar en README el requisito de un proxy TLS (nginx/Caddy) para despliegues publicos")

doc.add_paragraph()
heading(doc, "3.2 Mediano plazo (1-2 meses)", 2, "E65100")
bullet(doc, "Migrar cache y rate limiting a Redis para soportar multiples replicas/workers de forma consistente")
bullet(doc, "Configurar Uvicorn con --workers o pasar a Gunicorn+Uvicorn workers para paralelismo real bajo carga")
bullet(doc, "Anadir circuit breaker o backoff exponencial para VirusTotal/Safe Browsing/Fact Check ante cuota excedida")
bullet(doc, "Configurar CI (GitHub Actions): lint (ruff) + pytest en cada push/PR")
bullet(doc, "Decidir el destino de phishing_detector/: integrarlo a PhishingService o documentarlo formalmente como modulo de investigacion separado")

doc.add_paragraph()
heading(doc, "3.3 Largo plazo (siguiente iteracion mayor)", 2, "B71C1C")
bullet(doc, "Monitoreo de drift: comparar distribucion de features de produccion vs dataset de entrenamiento periodicamente")
bullet(doc, "Pipeline de reentrenamiento semi-automatico con validacion antes de promover un modelo nuevo")
bullet(doc, "Métricas de latencia por sub-servicio (WHOIS, HTML fetch, cada API externa, cada modelo ML) en vez de solo analysis_time_ms total")
bullet(doc, "Consolidar la cobertura de tests duplicada en una sola suite de referencia por modulo")

doc.add_paragraph()
heading(doc, "3.4 Matriz de priorizacion", 2, "1E3A5F")
add_table(doc,
    ["Accion", "Impacto", "Esfuerzo", "Plazo"],
    [
        ["Forzar API_KEY en produccion", "Alto", "Bajo", "Corto"],
        ["Healthcheck en compose", "Medio", "Bajo", "Corto"],
        ["Logging de requests", "Medio", "Bajo", "Corto"],
        ["Cache/rate-limit en Redis", "Alto", "Medio", "Mediano"],
        ["Multiples workers Uvicorn", "Alto", "Medio", "Mediano"],
        ["CI con lint + tests", "Medio", "Bajo", "Mediano"],
        ["Circuit breaker APIs externas", "Medio", "Medio", "Mediano"],
        ["Monitoreo de drift", "Medio", "Alto", "Largo"],
        ["Pipeline de reentrenamiento", "Medio", "Alto", "Largo"],
    ],
    col_widths=[5.5, 2.0, 2.0, 2.0],
)

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Diagnostico de Arquitectura: Componentes, Riesgos y Plan de Mejora  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "diagnostico_arquitectura_riesgos_plan_mejora.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
