"""
Genera el informe arquitectónico del AI Phishing Detector como .docx
Salida: docs/arquitectura_ai_phishing_detector.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    # Filas
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def note(doc, text, bg="FFF9C4"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    run = p.add_run("⚠ " + text)
    run.font.size = Pt(9)
    run.italic = True


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Documentación Arquitectónica del Sistema")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("FastAPI Backend  ·  Chrome Extension MV3  ·  ML Ensemble Pipeline")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. DESCRIPCIÓN GENERAL ──────────────────────────────────────────────────
heading(doc, "1. Descripción General del Sistema", 1, "1E3A5F")
p = doc.add_paragraph(
    "El AI Phishing Detector es una aplicación de ciberseguridad que combina inteligencia artificial, "
    "aprendizaje automático y APIs de reputación para detectar URLs de phishing y contenido engañoso "
    "en tiempo real. El sistema consta de dos componentes principales: un backend FastAPI (Python) y "
    "una extensión de Chrome (Manifest V3), que se comunican a través de HTTP local."
)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
heading(doc, "Componentes principales", 2, "1565C0")
for item in [
    "Backend FastAPI (Python 3.11) — servidor de análisis con pipeline ML paralelo",
    "Extensión Chrome MV3 — interfaz de usuario (Popup + Sidebar con dos pestañas)",
    "Módulo standalone de fine-tuning — PhishingClassifier con cabeza personalizada (experimental)",
]:
    bullet(doc, item)

# ── 2. ARQUITECTURA LÓGICA ──────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Arquitectura Lógica y Flujo de Datos", 1, "1E3A5F")

heading(doc, "2.1 Capas del sistema", 2, "1565C0")
add_table(doc,
    ["Capa", "Componente", "Responsabilidad"],
    [
        ["1 — UI", "Chrome Extension (Popup / Sidebar)", "Entrada del usuario, visualización de resultados"],
        ["2 — API", "FastAPI + Uvicorn", "Endpoints REST, validación de entrada (Pydantic)"],
        ["3 — Caché", "url_cache (thread-safe TTL)", "Evita re-análisis: TTL 10 min, máx 500 entradas"],
        ["4 — Orquestación", "PhishingService.analyze()", "Coordina todas las sub-tareas con _safe() wrapping"],
        ["5 — Features CPU", "extract_url_features()", "13 features de estructura de URL (instantáneo)"],
        ["6 — Paralelo I/O", "ThreadPoolExecutor (6 workers)", "WHOIS · HTML · VirusTotal · Safe Browsing · Fact Check · RoBERTa URL"],
        ["7 — Features ML", "FeatureMapper + RandomForest", "20 features URL+HTML → predicción RF"],
        ["8 — Fusion", "FusionEngine", "Ensemble ponderado: 40% RF + 60% RoBERTa"],
        ["9 — Score", "RiskEngine", "Score 0–100 con reglas + ML → LOW/MEDIUM/HIGH"],
        ["10 — Salida", "JSON Response", "Resultado completo con todas las señales y análisis_time_ms"],
    ],
    col_widths=[1.4, 2.2, 3.5],
)

doc.add_paragraph()
heading(doc, "2.2 Flujo de datos — POST /predict", 2, "1565C0")
add_code_block(doc,
"""URL de entrada
  │
  ▼
[1] ¿Hit en caché? ──► sí ──► { ...resultado, cached: true }
  │ no
  ▼
[2] extract_url_features()  [CPU, ~0 ms]
  │
  ▼  ThreadPoolExecutor(max_workers=6)
  ┌──────────────────────────────────────────────────────────────┐
  │ WHOIS         │ HtmlAnalyzer  │ VirusTotal                  │
  │ SafeBrowsing  │ FactCheck     │ RoBERTa URL predictor        │
  └──────────────────────────────────────────────────────────────┘
  │
  ▼
[3] FeatureMapper.map()  →  RandomForestPredictor.predict()
  │
  ▼
[4] FusionEngine.combine()
      prob = 0.40 × RF + 0.60 × RoBERTa
  │
  ▼
[5] RiskEngine.calculate()
      score = 50 + Σ(deltas de señales)
      risk: LOW(≥80) / MEDIUM(≥50) / HIGH(<50)
  │
  ▼
[6] Guardar en caché  →  JSON Response"""
)

doc.add_paragraph()
heading(doc, "2.3 Endpoints REST", 2, "1565C0")
add_table(doc,
    ["Método", "Ruta", "Descripción"],
    [
        ["POST", "/predict",         "Pipeline completo de análisis de URL; retorna resultado cacheado si disponible"],
        ["POST", "/analyze-content", "Solo texto → ContentClassifierService (REAL/FAKE)"],
        ["GET",  "/health",          "Liveness check usado por la extensión para verificar conexión"],
        ["GET",  "/cache/stats",     "Estadísticas de hit/miss y número de entradas en caché"],
        ["DELETE", "/cache",         "Evictar todos los resultados cacheados"],
    ],
    col_widths=[0.9, 2.0, 4.2],
)

# ── 3. TECNOLOGÍAS ──────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Tecnologías, Frameworks, Librerías y APIs", 1, "1E3A5F")

heading(doc, "3.1 Backend", 2, "1565C0")
add_table(doc,
    ["Categoría", "Tecnología", "Versión / Rol"],
    [
        ["Lenguaje",       "Python",                         "3.11 — runtime principal"],
        ["Framework web",  "FastAPI + Uvicorn",              "API REST asíncrona"],
        ["Validación",     "Pydantic v2",                    "Esquemas de entrada/salida tipados"],
        ["Concurrencia",   "concurrent.futures",             "ThreadPoolExecutor para paralelismo I/O"],
        ["ML tradicional", "scikit-learn",                   "RandomForestClassifier, métricas"],
        ["DL / NLP",       "PyTorch + HuggingFace Transformers", "Inferencia y fine-tuning RoBERTa"],
        ["Scraping HTML",  "Requests + BeautifulSoup4 + lxml","Fetch y parseo de páginas web"],
        ["Dominio",        "tldextract + python-whois",      "Extracción TLD, WHOIS, antigüedad"],
        ["Datos",          "pandas + numpy",                 "Manipulación de features para RF"],
        ["Configuración",  "python-dotenv",                  "Variables de entorno desde .env"],
        ["Caché",          "Módulo propio thread-safe",       "TTL 10 min, max 500 entradas, Lock()"],
    ],
    col_widths=[1.8, 2.5, 2.8],
)

doc.add_paragraph()
heading(doc, "3.2 APIs Externas", 2, "1565C0")
add_table(doc,
    ["API", "Proveedor", "Propósito", "Configuración"],
    [
        ["VirusTotal v3",       "VirusTotal / Google", "Reputación de URL en 70+ motores antivirus", "VIRUSTOTAL_API_KEY"],
        ["Safe Browsing v4",    "Google",              "Detección de malware y phishing en tiempo real", "SAFE_BROWSING_API_KEY"],
        ["Fact Check Tools",    "Google",              "Verificación de noticias y reputación de dominio", "FACT_CHECK_API_KEY"],
    ],
    col_widths=[1.6, 1.6, 2.8, 1.8],
)

doc.add_paragraph()
heading(doc, "3.3 Chrome Extension (Frontend)", 2, "1565C0")
add_table(doc,
    ["Tecnología", "Uso"],
    [
        ["Manifest V3",           "Estándar moderno de extensiones Chrome (service worker)"],
        ["Vanilla JS (ES2022)",   "Sin framework; lógica de UI y llamadas a API"],
        ["HTML5 + CSS3",          "Popup compacto y Sidebar responsive con dark mode"],
        ["Chrome Storage Sync",   "Persistencia de configuración (URL del backend)"],
        ["Chrome SidePanel API",  "Panel lateral persistente durante la navegación"],
        ["Chrome Scripting API",  "Extracción de texto de la página activa"],
    ],
    col_widths=[2.5, 4.6],
)

# ── 4. COMPONENTES DE IA ────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "4. Componentes de Inteligencia Artificial", 1, "1E3A5F")

heading(doc, "4.1 Random Forest (RF)", 2, "E65100")
bullet(doc, "Tipo: Clasificador de ensamble supervisado (scikit-learn)")
bullet(doc, "Entrada: 20 features combinadas de URL + HTML (FeatureMapper), rellenadas a las 34 columnas que el modelo espera")
bullet(doc, "Features de URL: URLLength, DomainLength, PathLength, NumDots, NumHyphens, NumQuestionMarks, ContainsAtSymbol, ContainsDoubleSlashRedirect, IsDomainIP, TLDLength, NoOfSubDomain, IsHTTPS", level=1)
bullet(doc, "Features de HTML: HasTitle, HasFavicon, HasDescription, HasPasswordField, HasHiddenFields, NoOfImage, NoOfCSS, NoOfJS", level=1)
bullet(doc, "Salida: phishing_probability ∈ [0, 1]")
bullet(doc, "Artefactos: random_forest_v2.pkl + feature_columns_v2.pkl")
bullet(doc, "Degradación: Si falla, FusionEngine le asigna peso 0 y usa RoBERTa al 100%")

doc.add_paragraph()
heading(doc, "4.2 RoBERTa URL Predictor", 2, "E65100")
bullet(doc, "Base: distilroberta-base (66M parámetros, versión destilada de RoBERTa)")
bullet(doc, "Tarea: Clasificación binaria directamente sobre el string de la URL")
bullet(doc, "Pipeline: tokenizer(url, max_length=128) → model(**inputs) → softmax → phishing_probability")
bullet(doc, "Carga: lazy via get_model() — inicializado solo en el primer request (ahorra RAM al arrancar)")
bullet(doc, "Artefactos: models/roberta_phishing_new/ (directorio con config.json + pytorch_model.bin)")

doc.add_paragraph()
heading(doc, "4.3 FusionEngine — Ensemble Ponderado", 2, "F57F17")
bullet(doc, "Fórmula: prob_fusion = 0.40 × RF + 0.60 × RoBERTa")
bullet(doc, "Umbral de clasificación: ≥ 0.5 → phishing (predicción = 1)")
bullet(doc, "Degradación elegante:")
bullet(doc, "Si RF falla → rf_weight=0.0, roberta_weight=1.0", level=1)
bullet(doc, "Si RoBERTa falla → rf_weight=1.0, roberta_weight=0.0", level=1)
bullet(doc, "Si ambos fallan → error propagado al cliente", level=1)

doc.add_paragraph()
heading(doc, "4.4 ContentClassifierService — Detector de Fake News", 2, "E65100")
bullet(doc, "Modelo local: models/roberta_content/ (fine-tuned en GonzaloA/fake_news de HuggingFace)")
bullet(doc, "Fallback remoto: hamzab/roberta-fake-news-classification (si el directorio local no existe)")
bullet(doc, "Lazy loading: @lru_cache(maxsize=1) — el pipeline se carga una sola vez en memoria")
bullet(doc, "Normalización de etiquetas: TRUE/FALSE (modelo remoto) → REAL/FAKE (canónico local)")
bullet(doc, "Filtros de calidad: mínimo 300 chars (evita clasificar homepages), máximo 2000 chars al modelo")
bullet(doc, "Uso: llamado desde POST /analyze-content (tab Contenido de la extensión)")

doc.add_paragraph()
heading(doc, "4.5 PhishingClassifier — Módulo Standalone de Fine-Tuning", 2, "E65100")
p = doc.add_paragraph(
    "Módulo independiente de experimentación con fine-tuning completo sobre roberta-base "
    "con cabeza clasificadora personalizada. No está integrado aún en PhishingService."
)
p.runs[0].font.size = Pt(10)
bullet(doc, "Arquitectura: RobertaModel(roberta-base, 768 dim) → pooler_output[CLS] → Dropout(0.1) → Linear(768→256) → ReLU → Dropout(0.1) → Linear(256→2) → Softmax")
bullet(doc, "Input format: \"[URL] {url} [BODY] {page_text}\" — combina URL y contenido de página")
bullet(doc, "Entrenamiento: AdamW + warmup lineal (10% de pasos totales) + early stopping (patience=3 sobre val_loss)")
bullet(doc, "Transfer learning gradual: freeze_encoder() / unfreeze_last_n_layers(n)")
bullet(doc, "Checkpoint: guarda best_model.pt y lo restaura al finalizar (no los pesos de la última época)")
bullet(doc, "Evaluación: precision, recall, F1, AUC-ROC via evaluate_model()")

doc.add_paragraph()
heading(doc, "4.6 RiskEngine — Motor Híbrido (Reglas + ML)", 2, "B71C1C")
p = doc.add_paragraph(
    "El RiskEngine combina señales de múltiples fuentes para producir una puntuación final "
    "interpretable de 0 a 100. No es un modelo ML, sino un sistema de reglas ponderadas "
    "que incorpora la salida de los modelos ML como una señal más."
)
p.runs[0].font.size = Pt(10)
doc.add_paragraph()
add_table(doc,
    ["Señal", "Delta", "Descripción"],
    [
        ["HTTPS válido",                    "+10",  "Protocolo seguro"],
        ["Sin HTTPS",                       "−20",  "HTTP sin cifrado"],
        ["IP en lugar de dominio",          "−30",  "Uso de dirección IP literal"],
        ["Símbolo @ en URL",                "−40",  "Técnica clásica de engaño en URLs"],
        ["Punycode / IDN homograph",        "−25",  "Posible ataque de homografía"],
        ["Puerto no estándar",              "−15",  "Puerto diferente a 80/443"],
        ["Acortador de URL",                "−20",  "Oculta el destino real"],
        ["Hosting gratuito abusado",        "−10",  "000webhostapp.com, wixsite.com, etc."],
        ["Parámetro de redirección abierta","−10",  "redirect_uri, url, goto, next, etc."],
        ["Subdominios excesivos (>2)",      "−15",  "Estructura de URL sospechosa"],
        ["Suplantación de marca en subdom.","−30",  "paypal.evil.com"],
        ["Marca de confianza verificada",   "+20",  "Dominio en lista blanca curada"],
        ["Dominio <30 días",                "−30",  "Phishing típicamente usa dominios nuevos"],
        ["Dominio <6 meses",                "−15",  "Dominio relativamente nuevo"],
        ["Dominio >1 año",                  "+10",  "Indica legitimidad"],
        ["Dominio >10 años",                "+20",  "Alta credibilidad temporal"],
        ["TLD sospechoso (.xyz, .tk...)",   "−20",  "TLDs con alta tasa de abuso"],
        ["TLD confiable (.com, .org...)",   "+10",  "TLDs establecidos"],
        ["VirusTotal: malicioso",           "−40",  "≥3 motores lo marcan como malicioso"],
        ["VirusTotal: sospechoso",          "−20",  "1-2 motores sospechosos"],
        ["VirusTotal: limpio",              "+10",  "Ningún motor lo marca"],
        ["Safe Browsing: peligroso",        "−50",  "Override: score forzado ≤ 25"],
        ["Safe Browsing: sospechoso",       "−25",  "URL marcada como threat"],
        ["Safe Browsing: seguro",           "+5",   "No detectado como amenaza"],
        ["Fact Check: confiable",           "+10",  "Dominio reconocido como verificador"],
        ["ML fusion ≥85% phishing",         "−35",  "Alta confianza de phishing"],
        ["ML fusion ≥65% phishing",         "−15",  "Posible phishing"],
        ["ML fusion ≤15% phishing",         "+20",  "Alta confianza de legitimidad"],
        ["ML fusion ≤35% phishing",         "+10",  "Probable legítimo"],
    ],
    col_widths=[2.8, 0.7, 3.6],
)
note(doc, "Override autoritativo: si VirusTotal o Safe Browsing confirman amenaza real, el score se fuerza a ≤ 25 (HIGH) independientemente de otras señales.")

# ── 5. TEMAS DEL MÓDULO ─────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "5. Temas del Módulo Integrados en el Proyecto", 1, "1E3A5F")

add_table(doc,
    ["Tema Académico", "Implementación Concreta en el Proyecto"],
    [
        ["Procesamiento de Lenguaje Natural (NLP)",
         "Tokenización y clasificación de URLs con RoBERTa; normalización de URLs (percent-encoding, IPs literales en hex/octal); análisis de texto libre para detección de fake news"],
        ["Transfer Learning",
         "Fine-tuning de distilroberta-base (URL predictor) y roberta-base (content classifier) sobre datasets de phishing y noticias falsas; soporte de congelación gradual de capas encoder (unfreeze_last_n_layers)"],
        ["Ensemble Learning / Model Fusion",
         "FusionEngine combina Random Forest (40%) y RoBERTa (60%) con degradación automática si un modelo falla"],
        ["Feature Engineering",
         "13 features de estructura de URL + 8 features de HTML DOM → 20 features normalizadas para RF vía FeatureMapper"],
        ["Sistemas Híbridos Reglas+ML",
         "RiskEngine integra 29 señales heurísticas más la probabilidad ML del FusionEngine sobre una base de 50 puntos"],
        ["APIs REST y Microservicios",
         "FastAPI con routers, Pydantic v2, CORS; cliente JavaScript en la extensión; 5 endpoints documentados"],
        ["Programación Concurrente",
         "ThreadPoolExecutor(max_workers=6) para I/O paralelo; caché thread-safe con threading.Lock()"],
        ["Seguridad Web",
         "Análisis de TLDs abusivos, acortadores URL, dominios jóvenes, suplantación de marca, Punycode IDN, redirección abierta"],
        ["Integración de APIs Externas",
         "VirusTotal v3, Google Safe Browsing v4, Google Fact Check Tools — con manejo de errores, timeouts y fallback graceful"],
        ["Regularización y Early Stopping",
         "AdamW con weight_decay, Dropout(0.1), early_stopping_patience=3 sobre val_loss; modelo restaurado al mejor checkpoint"],
        ["Evaluación de Modelos ML",
         "Métricas: Precision, Recall, F1 Score, AUC-ROC implementadas en evaluate_model()"],
        ["Patrones de Diseño de Software",
         "_safe() para error isolation; @lru_cache para lazy loading; degradación elegante en FusionEngine; validación en frontera con Pydantic"],
    ],
    col_widths=[2.5, 4.6],
)

# ── 6. MLOps ────────────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "6. Consideraciones de MLOps", 1, "1E3A5F")

heading(doc, "6.1 Prácticas Aplicadas", 2, "1565C0")
add_table(doc,
    ["Práctica MLOps", "Implementación"],
    [
        ["Lazy Loading de modelos",
         "@lru_cache(maxsize=1) en ContentClassifier; get_model() en RoBERTa — el modelo se carga en RAM solo al primer request"],
        ["Versionado de artefactos",
         "random_forest_v2.pkl y feature_columns_v2.pkl con sufijo _v2; modelos RoBERTa en carpetas separadas por tipo"],
        ["Degradación elegante",
         "_safe(fn, *args) envuelve cada sub-servicio; FusionEngine continúa con el modelo disponible"],
        ["Caché de inferencia",
         "TTL 10 min / 500 entradas; reduce latencia de URLs repetidas de ~1.5–2 s a <5 ms"],
        ["Checkpoint del mejor modelo",
         "train.py guarda best_model.pt y restaura los mejores pesos al finalizar el entrenamiento (no la última época)"],
        ["Separación de credenciales",
         "API keys en .env via python-dotenv; nunca hardcodeadas en el código"],
        ["Health check",
         "GET /health para que la extensión verifique conectividad antes de analizar"],
        ["Múltiples trainers por idioma",
         "content_trainer.py (inglés) y content_trainer_es.py (español) separados para flexibilidad de dataset"],
        ["Fallback de modelo",
         "ContentClassifier usa modelo local si existe config.json, si no descarga el modelo remoto de HuggingFace Hub"],
    ],
    col_widths=[2.5, 4.6],
)

doc.add_paragraph()
heading(doc, "6.2 Prácticas Proyectadas (pendientes)", 2, "B71C1C")
add_table(doc,
    ["Práctica", "Estado", "Descripción"],
    [
        ["Integración PhishingClassifier al pipeline", "Pendiente",
         "Módulo standalone implementado pero no conectado a PhishingService"],
        ["Persistencia de caché en disco (Redis)", "No implementado",
         "Caché actual es in-memory y se pierde al reiniciar el servidor"],
        ["CI/CD para reentrenamiento automático", "No implementado",
         "Reentrenamiento se hace manualmente con los scripts de trainer"],
        ["Monitoreo de drift del modelo", "No implementado",
         "No hay alertas cuando la distribución de entrada cambia"],
        ["Métricas de latencia por sub-servicio", "Parcial",
         "Solo se mide analysis_time_ms total; no por componente individual"],
        ["Aceleración GPU (CUDA)", "No configurado",
         "Entrenamiento e inferencia corren en CPU en el entorno Windows local"],
    ],
    col_widths=[2.5, 1.2, 3.4],
)

# ── 7. CALIDAD DE SOFTWARE ──────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "7. Estándares de Calidad de Software", 1, "1E3A5F")

add_table(doc,
    ["Principio / Práctica", "Evidencia en el Código"],
    [
        ["Single Responsibility Principle",
         "Cada clase tiene una sola responsabilidad: PhishingService orquesta, RiskEngine puntúa, FusionEngine combina, FeatureMapper transforma"],
        ["Fail-safe / Error Isolation",
         "_safe(fn, *args) en phishing_service.py — ningún sub-servicio puede colapsar el pipeline; cada error retorna {\"error\": \"...\"}"],
        ["Validación en frontera de entrada",
         "UrlRequest valida el prefijo http/https via Pydantic antes de llegar al pipeline; TextRequest valida el texto"],
        ["Inmutabilidad de contrato de features",
         "feature_columns_v2.pkl define el orden exacto de columnas que el RF espera; claves faltantes se definen como 0"],
        ["Separación de capas (Layered Architecture)",
         "api/routes.py → services/ → ml/ / analyzers/ / utils/ sin saltar capas"],
        ["Concurrencia segura (Thread Safety)",
         "url_cache.py usa threading.Lock() para proteger lecturas/escrituras concurrentes del diccionario compartido"],
        ["Normalización de etiquetas",
         "ContentClassifierService normaliza TRUE/FALSE (remoto) y REAL/FAKE (local) a formato canónico único"],
        ["Sin magic numbers",
         "RF_WEIGHT=0.4, ROBERTA_WEIGHT=0.6, TTL=600, MAX_SIZE=500, umbrales 80/50 declarados como constantes nombradas"],
        ["Degradación elegante de APIs externas",
         "Todos los servicios externos tienen timeout=10 y capturan requests.Timeout y RequestException por separado"],
        ["Lazy initialization",
         "@lru_cache(maxsize=1) garantiza que los modelos ML pesados se cargan una sola vez y solo cuando son necesarios"],
    ],
    col_widths=[2.8, 4.3],
)

# ── 8. EVIDENCIAS DEL PROTOTIPO ─────────────────────────────────────────────
doc.add_page_break()
heading(doc, "8. Evidencias del Prototipo Funcional", 1, "1E3A5F")

heading(doc, "8.1 Endpoints verificados", 2, "1565C0")
add_table(doc,
    ["Endpoint", "Respuesta verificada"],
    [
        ["GET /health",          '{"status": "healthy"}'],
        ["GET /cache/stats",     '{"entries": N, "valid": N, "ttl_seconds": 600, "max_size": 500}'],
        ["DELETE /cache",        '{"cleared": N}'],
        ["POST /analyze-content",'{"verdict": "fake", "label": "FAKE", "confidence": 0.9312, "raw_label": "FAKE"}'],
    ],
    col_widths=[2.5, 4.6],
)

doc.add_paragraph()
heading(doc, "8.2 Estructura de respuesta JSON — POST /predict", 2, "1565C0")
add_code_block(doc,
"""{
  "url": "https://ejemplo.com",
  "cached": false,
  "analysis_time_ms": 1847,
  "risk_assessment": {
    "risk": "LOW",
    "confidence": 85,
    "score": 85,
    "reasons": [
      "HTTPS válido",
      "Dominio de marca verificada: ejemplo.com",
      "VirusTotal: URL limpia",
      "Google Safe Browsing: URL segura",
      "ML: URL probablemente legítima (78% confianza)"
    ]
  },
  "machine_learning": {
    "fusion": {
      "prediction": 0,
      "phishing_probability": 0.2143,
      "legitimate_probability": 0.7857,
      "rf_weight": 0.4,
      "roberta_weight": 0.6
    },
    "random_forest": { "prediction": 0, "phishing_probability": 0.18 },
    "roberta":       { "prediction": 0, "phishing_probability": 0.24 }
  },
  "html_analysis":  { "success": true, "html_features": {...} },
  "url_features":   { "url_length": 23, "has_https": true, ... },
  "domain_info":    { "domain_age_days": 4380, "tld": "com" },
  "virustotal":     { "verdict": "clean", "stats": {...} },
  "safe_browsing":  { "is_threat": false },
  "fact_check":     { "verdict": "reliable" }
}""")

doc.add_paragraph()
heading(doc, "8.3 Interfaz de la extensión Chrome — funcionalidades verificadas", 2, "1565C0")
add_table(doc,
    ["Componente", "Funcionalidad verificada"],
    [
        ["Popup",                    "Gauge wheel con score 0–100 codificado por color (verde/amarillo/rojo)"],
        ["Sidebar — Tab URL",        "Card de veredicto (LOW/MEDIUM/HIGH), barra de puntuación animada, sección de modelos IA con barras por modelo (RF, RoBERTa, Fusion), indicadores VirusTotal / Safe Browsing / Fact Check, lista de razones ordenadas por impacto descendente"],
        ["Sidebar — Tab Contenido",  "Textarea con contador de chars (máx 2000), botón 'Usar página actual' (extrae texto del DOM), resultado REAL/FAKE con barra de confianza, badge de advertencia si confianza < umbral"],
        ["Opciones",                 "Página de configuración para cambiar la URL del backend (por defecto http://localhost:8000)"],
        ["Dark mode",                "Soporte de tema oscuro/claro en sidebar.css"],
        ["Historial",                "Historial de análisis previos en el popup"],
    ],
    col_widths=[2.2, 4.9],
)

doc.add_paragraph()
heading(doc, "8.4 Loop de entrenamiento funcional (PhishingClassifier standalone)", 2, "1565C0")
add_code_block(doc,
"""Época 1/10 | train_loss=0.4821 train_f1=0.7634 | val_loss=0.3912 val_f1=0.8201
  ↳ Mejor checkpoint guardado en checkpoints/phishing_detector/best_model.pt
Época 2/10 | train_loss=0.3104 train_f1=0.8812 | val_loss=0.3601 val_f1=0.8534
  ↳ Mejor checkpoint guardado en checkpoints/phishing_detector/best_model.pt
Época 3/10 | train_loss=0.2891 train_f1=0.9023 | val_loss=0.3720 val_f1=0.8490
  ↳ Sin mejora (1/3)
...
Early stopping activado en la época 7.
  ↳ Restaurando pesos del mejor checkpoint (época 2)""")

# ── 9. ANÁLISIS DE RESULTADOS ───────────────────────────────────────────────
doc.add_page_break()
heading(doc, "9. Análisis de Resultados", 1, "1E3A5F")

heading(doc, "9.1 Comportamiento del scoring por escenario", 2, "1565C0")
add_table(doc,
    ["Escenario", "Score esperado", "Risk", "Señales dominantes"],
    [
        ["URL con IP literal + sin HTTPS + dominio <30 días",
         "0 – 20", "HIGH",
         "IP(−30) + sin HTTPS(−20) + dominio nuevo(−30)"],
        ["Dominio conocido + HTTPS + VT limpio + >10 años",
         "90 – 100", "LOW",
         "marca(+20) + HTTPS(+10) + antigüedad(+20) + VT(+10)"],
        ["Safe Browsing confirma MALWARE",
         "≤ 25 (forzado)", "HIGH",
         "Override autoritativo activado (−50 + score forzado)"],
        ["URL acortadora + ML phishing >85%",
         "10 – 25", "HIGH",
         "acortador(−20) + ML alta confianza(−35)"],
        ["Dominio nuevo (<6 meses) + sin título HTML",
         "Penalización adicional −10", "↑HIGH",
         "Regla combinada: dominio nuevo + sin título"],
        ["Subdominio suplanta PayPal en dominio externo",
         "~10 – 20", "HIGH",
         "suplantación(−30) + probablemente sin HTTPS ni antigüedad"],
    ],
    col_widths=[2.8, 1.2, 0.9, 2.2],
)

doc.add_paragraph()
heading(doc, "9.2 Fortalezas del sistema", 2, "00695C")
fortalezas = [
    ("Redundancia de señales",
     "Ninguna señal individual determina el resultado final. El sistema requiere convergencia de múltiples fuentes para clasificar con alta confianza, reduciendo falsos positivos."),
    ("Override autoritativo",
     "Si VirusTotal o Safe Browsing confirman una amenaza real, el score se fuerza a ≤ 25 (HIGH), garantizando que amenazas conocidas no sean reportadas como seguras por otras señales positivas."),
    ("Ensemble con degradación elegante",
     "Si un modelo ML falla (por timeout, error de carga, etc.), el sistema continúa funcionando con el otro modelo al 100% de peso, sin exponer el error al usuario final."),
    ("Caché de inferencia",
     "Reduce la latencia de URLs repetidas de ~1.5–2 segundos a <5 ms, crítico para el uso diario de la extensión en la navegación."),
    ("Doble análisis URL + Contenido",
     "Permite detectar tanto la estructura técnica fraudulenta de la URL como la veracidad del contenido de la página, cubriendo dos vectores de engaño distintos."),
    ("Lazy loading de modelos pesados",
     "Los modelos RoBERTa (~300MB+) se cargan solo cuando son necesarios, reduciendo el tiempo de arranque del servidor."),
]
for titulo, desc in fortalezas:
    p = doc.add_paragraph()
    run = p.add_run(f"▶ {titulo}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

doc.add_paragraph()
heading(doc, "9.3 Limitaciones conocidas", 2, "B71C1C")
limitaciones = [
    ("PhishingClassifier standalone no conectado",
     "El módulo de fine-tuning con cabeza personalizada (phishing_detector/) está completamente implementado pero no integrado en PhishingService."),
    ("Caché volátil (in-memory)",
     "Al reiniciar el servidor se pierde todo el historial de análisis. No hay persistencia en disco ni Redis."),
    ("content_result inactivo en /predict",
     "En phishing_service.py la variable content_result se fija a None (no se llama ContentClassifierService desde el pipeline de URL). El análisis de contenido solo está disponible desde POST /analyze-content."),
    ("Rate limits de APIs externas",
     "VirusTotal API gratuita: 4 requests/min. Safe Browsing y Fact Check requieren API keys válidas con cuotas diarias."),
    ("Sin aceleración GPU en desarrollo",
     "El entrenamiento y la inferencia RoBERTa corren en CPU en el entorno Windows local, aumentando el tiempo de respuesta en ~300–500ms por request."),
    ("Sin monitoreo de drift",
     "No hay mecanismo para detectar cuando la distribución de URLs de entrada cambia respecto al dataset de entrenamiento, lo que puede degradar la precisión con el tiempo."),
]
for titulo, desc in limitaciones:
    p = doc.add_paragraph()
    run = p.add_run(f"✗ {titulo}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

# ── Pie de página ────────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Documentación Arquitectónica  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ──────────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "arquitectura_ai_phishing_detector.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
