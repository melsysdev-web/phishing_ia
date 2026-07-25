"""
Genera el informe final y completo: "API de IA testeable desde un cliente
externo" del AI Phishing Detector como .docx. Consolida en un solo documento
todo el producto esperado del requisito (codigo, endpoints, contratos,
validacion, manejo de errores, evidencia de prueba, README y docs/api.md).
Salida: docs/informe_final_api_ia.docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def para(doc, text, size=9.5, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Informe Final — API de IA Testeable desde un Cliente Externo")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Codigo · Endpoints · Contratos · Validacion · Errores · Pruebas · Documentacion")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 0. TABLA DE CONTENIDO (manual) ────────────────────────────────────────
heading(doc, "Contenido", 1, "1E3A5F")
toc_items = [
    "1. Requisito",
    "2. Resumen Ejecutivo — Producto Esperado",
    "3. Codigo Actualizado del Proyecto",
    "4. Endpoint Principal de IA (POST /predict, POST /analyze-content)",
    "5. Endpoint /health",
    "6. Endpoint /metadata",
    "7. Contrato de Entrada y Salida",
    "8. Validacion Basica de Datos",
    "9. Manejo Controlado de Errores",
    "10. Evidencia de Prueba",
    "11. README Actualizado",
    "12. Documento docs/api.md",
    "13. Como Probar la API desde un Cliente Externo",
    "14. Conclusion",
]
for item in toc_items:
    bullet(doc, item)

# ── 1. REQUISITO ─────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "1. Requisito", 1, "1E3A5F")
p = doc.add_paragraph(
    "“La API debe permitir que la funcionalidad de IA pueda ser probada desde "
    "un cliente externo como Swagger, curl, Postman o una interfaz web simple.”"
)
p.runs[0].italic = True
p.runs[0].font.size = Pt(11)

# ── 2. RESUMEN EJECUTIVO ──────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Resumen Ejecutivo — Producto Esperado", 1, "1E3A5F")
p = doc.add_paragraph(
    "Estado de cada entregable exigido, verificado contra el codigo, la suite "
    "de tests y la documentacion actuales del repositorio (rama "
    "phishing-ia/risk-engine-v2). Los 10 puntos estan cerrados."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["#", "Entregable", "Estado", "Seccion"],
    [
        ["1", "Codigo actualizado del proyecto", "Cumple", "3"],
        ["2", "Endpoint principal de IA", "Cumple", "4"],
        ["3", "Endpoint /health", "Cumple", "5"],
        ["4", "Endpoint /metadata o equivalente", "Cumple", "6"],
        ["5", "Contrato de entrada y salida", "Cumple", "7"],
        ["6", "Validacion basica de datos", "Cumple", "8"],
        ["7", "Manejo controlado de errores", "Cumple", "9"],
        ["8", "Evidencia de prueba", "Cumple", "10"],
        ["9", "README actualizado", "Cumple", "11"],
        ["10", "Documento docs/api.md o equivalente", "Cumple", "12"],
    ],
    col_widths=[0.8, 6.5, 2.0, 1.5],
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Resultado global: ")
run.bold = True
run.font.size = Pt(10)
run2 = p.add_run("268/268 tests aprobados. La API es accesible y verificable desde Swagger UI, curl y Postman sin necesidad de leer el codigo fuente.")
run2.font.size = Pt(10)

# ── 3. CODIGO ACTUALIZADO ─────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Codigo Actualizado del Proyecto", 1, "1E3A5F")
add_table(doc,
    ["Archivo", "Cambio"],
    [
        ["backend/app/main.py", "GET /metadata; exception handler global (@app.exception_handler(Exception)); tags/response_model en / y /health; openapi_tags"],
        ["backend/app/api/routes.py", "response_model, tags, summary, description en /predict, /analyze-content, /cache/stats, DELETE /cache; responses={500: ErrorResponse} en /predict y /analyze-content"],
        ["backend/app/schemas/response_schema.py", "Nuevo: PredictResponse, ContentAnalysisResponse, CacheStatsResponse, CacheClearResponse, HealthResponse, RootResponse, MetadataResponse, ModelsAvailability, ErrorResponse"],
        ["backend/app/schemas/request_schema.py", "Field(description=...) en UrlRequest.url y TextRequest.text"],
        ["tests/test_cache_and_security.py", "Nuevo: 11 tests — /cache/stats, DELETE /cache, /metadata, 403, 429, 500 controlado"],
        ["tests/test_api.py", "Fixtures mock_safe_response/mock_phishing_response completados para cumplir el contrato de PredictResponse"],
        ["tests/conftest.py", "Fixture autouse que resetea el rate limiter antes de cada test (aislamiento entre archivos)"],
        ["README.md", "Seccion 'Probar la API sin escribir codigo'; tabla de endpoints con /health y /metadata; link a docs/api.md"],
        ["docs/api.md", "Nuevo: referencia completa de la API"],
    ],
    col_widths=[4.5, 9.5],
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Base: ")
run.bold = True
run.font.size = Pt(9.5)
run2 = p.add_run("commit cbc3fd5 (\"feat: Docker deployment, API auth/rate-limiting, standalone fine-tuning module, and cleanup\"), rama phishing-ia/risk-engine-v2. Los cambios de esta sesion estan en el arbol de trabajo, pendientes de commit.")
run2.font.size = Pt(9.5)

# ── 4. ENDPOINT PRINCIPAL DE IA ───────────────────────────────────────────
doc.add_page_break()
heading(doc, "4. Endpoint Principal de IA", 1, "1E3A5F")

heading(doc, "4.1 POST /predict", 2, "1565C0")
p = doc.add_paragraph(
    "Pipeline completo de deteccion de phishing: WHOIS, HTML, VirusTotal, "
    "Google Safe Browsing, Google Fact Check, Random Forest y RoBERTa, "
    "fusionados en un score 0-100 con razones explicables. Resultado "
    "cacheado 10 minutos (TTL) por URL."
)
p.runs[0].font.size = Pt(9.5)
add_code_block(doc,
"""curl -X POST http://localhost:8000/predict \\
  -H "Content-Type: application/json" \\
  -H "X-API-Key: <tu_api_key>" \\
  -d '{"url": "https://ejemplo.com"}'""")
add_code_block(doc,
"""// Response 200
{
  "url": "https://ejemplo.com",
  "cached": false,
  "analysis_time_ms": 1847,
  "risk_assessment": { "risk": "LOW", "confidence": 85, "score": 85, "reasons": [...] },
  "machine_learning": {
    "fusion":        { "prediction": 0, "phishing_probability": 0.21, "legitimate_probability": 0.79, "rf_weight": 0.4, "roberta_weight": 0.6 },
    "random_forest": { "prediction": 0, "phishing_probability": 0.18, "legitimate_probability": 0.82 },
    "roberta":       { "prediction": 0, "phishing_probability": 0.24, "legitimate_probability": 0.76 }
  },
  "html_analysis": {...}, "url_features": {...}, "domain_info": {...},
  "virustotal": {...}, "safe_browsing": {...}, "fact_check": {...},
  "content_classification": null
}""")

doc.add_paragraph()
heading(doc, "4.2 POST /analyze-content", 2, "1565C0")
p = doc.add_paragraph(
    "Clasifica texto libre como REAL o FAKE (fake news / contenido enganoso). "
    "Textos con menos de 300 caracteres devuelven verdict='no_content' sin "
    "ejecutar el modelo."
)
p.runs[0].font.size = Pt(9.5)
add_code_block(doc,
"""curl -X POST http://localhost:8000/analyze-content \\
  -H "Content-Type: application/json" \\
  -H "X-API-Key: <tu_api_key>" \\
  -d '{"text": "un articulo largo de al menos 300 caracteres..."}'

// Response 200
{ "verdict": "real", "label": "REAL", "confidence": 0.93, "raw_label": "REAL" }""")

# ── 5. ENDPOINT /health ───────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "5. Endpoint /health", 1, "1E3A5F")
p = doc.add_paragraph(
    "Liveness check sin autenticacion. Usado por la extension de Chrome para "
    "verificar conexion y por el HEALTHCHECK de Docker Compose."
)
p.runs[0].font.size = Pt(9.5)
add_code_block(doc, 'curl http://localhost:8000/health\n\n// Response 200\n{ "status": "healthy" }')

# ── 6. ENDPOINT /metadata ─────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "6. Endpoint /metadata", 1, "1E3A5F")
p = doc.add_paragraph(
    "Expone version de la API, disponibilidad de cada modelo ML y "
    "configuracion activa (rate limit, cache), sin autenticacion y sin cargar "
    "los modelos en memoria (solo verifica existencia de archivos via "
    "get_models_dir())."
)
p.runs[0].font.size = Pt(9.5)
add_code_block(doc,
"""curl http://localhost:8000/metadata

// Response 200
{
  "api_version": "1.0.0",
  "models": { "random_forest": true, "roberta_url": true, "roberta_content": true },
  "rate_limit_per_minute": 30,
  "cache_ttl_seconds": 600,
  "cache_max_size": 500
}""")
bullet(doc, "models.* es false si el artefacto correspondiente no existe en models/ (o MODELS_DIR en despliegue) — permite a un cliente externo verificar capacidades antes de llamar a /predict")

# ── 7. CONTRATO DE ENTRADA Y SALIDA ───────────────────────────────────────
doc.add_page_break()
heading(doc, "7. Contrato de Entrada y Salida", 1, "1E3A5F")

heading(doc, "7.1 Schemas de entrada (backend/app/schemas/request_schema.py)", 2, "1565C0")
add_table(doc,
    ["Schema", "Campo", "Tipo", "Regla"],
    [
        ["UrlRequest", "url", "str", "Debe comenzar con http:// o https:// (field_validator); si no, 422"],
        ["TextRequest", "text", "str", "Sin minimo a nivel de schema; < 300 caracteres se resuelve como no_content en el servicio, no como error de validacion"],
    ],
    col_widths=[2.5, 1.8, 1.5, 8.2],
)

doc.add_paragraph()
heading(doc, "7.2 Schemas de salida (backend/app/schemas/response_schema.py)", 2, "1565C0")
add_table(doc,
    ["Schema", "Usado en", "Notas"],
    [
        ["PredictResponse", "POST /predict", "Top-level tipado (url, cached, analysis_time_ms, content_classification); sub-diccionarios como dict abierto porque pueden degradarse a {\"error\": \"...\"}"],
        ["ContentAnalysisResponse", "POST /analyze-content", "Todos los campos opcionales: cubre las 3 formas validas (clasificado / no_content / error)"],
        ["CacheStatsResponse", "GET /cache/stats", "entries, valid, ttl_seconds, max_size — tipado estricto (forma siempre estable)"],
        ["CacheClearResponse", "DELETE /cache", "cleared: int"],
        ["HealthResponse", "GET /health", "status: str"],
        ["RootResponse", "GET /", "message: str"],
        ["MetadataResponse", "GET /metadata", "api_version, models (ModelsAvailability), rate_limit_per_minute, cache_ttl_seconds, cache_max_size"],
        ["ErrorResponse", "Respuestas 500 documentadas", "error: str, detail: Optional[str]"],
    ],
    col_widths=[3.3, 3.3, 6.9],
)

# ── 8. VALIDACION BASICA DE DATOS ─────────────────────────────────────────
doc.add_page_break()
heading(doc, "8. Validacion Basica de Datos", 1, "1E3A5F")
bullet(doc, "Pydantic valida tipos automaticamente (url: str, text: str) — un body con tipo incorrecto o campo faltante devuelve 422 sin llegar al handler")
bullet(doc, "UrlRequest.validate_url (field_validator): exige prefijo http:// o https://, sino ValueError -> 422")
bullet(doc, "El formato de error 422 es el estandar de FastAPI: {\"detail\": [{\"loc\": [...], \"msg\": \"...\", \"type\": \"...\"}]}")
add_code_block(doc,
"""curl -X POST http://localhost:8000/predict -d '{"url": "no-es-una-url"}'

-> 422
{
  "detail": [
    { "loc": ["body","url"], "msg": "Value error, La URL debe comenzar con http:// o https://", "type": "value_error" }
  ]
}""")

# ── 9. MANEJO CONTROLADO DE ERRORES ───────────────────────────────────────
doc.add_page_break()
heading(doc, "9. Manejo Controlado de Errores", 1, "1E3A5F")
p = doc.add_paragraph("Estrategia de manejo de errores en 3 capas:")
p.runs[0].font.size = Pt(9.5)

heading(doc, "9.1 Nivel sub-servicio (degradacion elegante)", 2, "1565C0")
bullet(doc, "APIs externas (VirusTotal, Safe Browsing, Fact Check): try/except requests.Timeout / requests.RequestException por servicio, degradan a {\"error\": \"...\"} dentro de su propia clave del response")
bullet(doc, "Sub-tareas del pipeline (WHOIS, HTML, modelos ML): envueltas en _safe(fn, *args) en phishing_service.py, mismo patron")
bullet(doc, "ContentClassifierService: try/except propio, retorna {\"error\": \"...\"} con status 200 (el error queda en el body)")
bullet(doc, "FusionEngine: si un modelo ML falla, usa el otro al 100% de peso en vez de fallar")

heading(doc, "9.2 Nivel HTTP (contratos explicitos)", 2, "1565C0")
add_table(doc,
    ["Codigo", "Cuando", "Contrato"],
    [
        ["422", "Body invalido / campo faltante", "{\"detail\": [{\"loc\": [...], \"msg\": \"...\", \"type\": \"...\"}]}"],
        ["403", "API_KEY configurada y X-API-Key falta/no coincide", "{\"detail\": \"API key invalida o ausente\"}"],
        ["429", ">30 solicitudes/60s a /predict o /analyze-content por IP", "{\"error\": \"Demasiadas solicitudes...\"}, header Retry-After: 60"],
        ["500", "Excepcion no prevista fuera de las rutas con manejo controlado", "{\"error\": \"Error interno del servidor\", \"detail\": \"...\"} — detail=null en produccion"],
    ],
    col_widths=[1.5, 6.0, 6.3],
)

heading(doc, "9.3 Nivel global (nuevo)", 2, "2E7D32")
bullet(doc, "@app.exception_handler(Exception) en main.py: red de seguridad final — cualquier excepcion que escape de las capas anteriores devuelve un 500 con contrato JSON consistente en vez del 500 generico de FastAPI")
bullet(doc, "Registra la excepcion via logger.exception() antes de responder (auditoria)")

# ── 10. EVIDENCIA DE PRUEBA ───────────────────────────────────────────────
doc.add_page_break()
heading(doc, "10. Evidencia de Prueba", 1, "1E3A5F")
p = doc.add_paragraph(
    "Los endpoints se prueban con FastAPI TestClient, sin depender de red ni "
    "modelos reales (loaders y servicios externos mockeados via conftest.py "
    "y unittest.mock.patch)."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Endpoint / mecanismo", "Archivo de tests", "# tests"],
    [
        ["GET /health, GET /", "tests/test_api.py", "4"],
        ["POST /predict", "tests/test_api.py", "12"],
        ["POST /analyze-content", "tests/test_content_api.py", "16"],
        ["GET /metadata", "tests/test_cache_and_security.py", "2"],
        ["GET /cache/stats, DELETE /cache", "tests/test_cache_and_security.py", "4"],
        ["require_api_key (403)", "tests/test_cache_and_security.py", "3"],
        ["RateLimitMiddleware (429)", "tests/test_cache_and_security.py", "1"],
        ["Exception handler global (500)", "tests/test_cache_and_security.py", "1"],
    ],
    col_widths=[5.0, 5.5, 2.3],
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Total: ")
run.bold = True
run.font.size = Pt(10)
run2 = p.add_run("268 tests, 268 aprobados. Comando: venv\\Scripts\\python -m pytest (pytest.ini: testpaths = tests backend/tests).")
run2.font.size = Pt(10)

# ── 11. README ACTUALIZADO ────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "11. README Actualizado", 1, "1E3A5F")
bullet(doc, "Tabla de endpoints ampliada con /health y /metadata")
bullet(doc, "Nueva seccion 'Probar la API sin escribir codigo': link a Swagger UI (/docs), instrucciones de import a Postman via /openapi.json, ejemplo curl completo")
bullet(doc, "Nueva fila en la tabla de Documentacion apuntando a docs/api.md")
bullet(doc, "Comando de pruebas corregido (referenciaba un archivo eliminado en una limpieza previa)")

# ── 12. docs/api.md ────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "12. Documento docs/api.md", 1, "1E3A5F")
p = doc.add_paragraph(
    "Referencia de API versionada en Markdown (vive junto al codigo en git, "
    "a diferencia de los informes .docx de esta sesion). Contenido:"
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "Tabla de los 7 endpoints (metodo, tag, auth, rate limit)")
bullet(doc, "Contrato detallado de cada endpoint con ejemplos curl y responses reales")
bullet(doc, "Tabla de codigos de error (422 / 403 / 429 / 500)")
bullet(doc, "Instrucciones para probar sin escribir codigo (Swagger UI, import a Postman)")

# ── 13. COMO PROBAR DESDE UN CLIENTE EXTERNO ──────────────────────────────
doc.add_page_break()
heading(doc, "13. Como Probar la API desde un Cliente Externo", 1, "1E3A5F")

heading(doc, "13.1 Swagger UI", 2, "1565C0")
bullet(doc, "Levantar el backend: venv\\Scripts\\uvicorn backend.app.main:app --reload")
bullet(doc, "Abrir http://localhost:8000/docs")
bullet(doc, "Expandir cualquier endpoint -> Try it out -> completar el body -> Execute")
bullet(doc, "Los 7 endpoints aparecen agrupados en 3 tags: Sistema, Analisis, Cache")

doc.add_paragraph()
heading(doc, "13.2 curl", 2, "1565C0")
add_code_block(doc,
"""curl http://localhost:8000/health
curl http://localhost:8000/metadata
curl -X POST http://localhost:8000/predict -H "Content-Type: application/json" -d '{"url": "https://ejemplo.com"}'""")

doc.add_paragraph()
heading(doc, "13.3 Postman", 2, "1565C0")
bullet(doc, "Import -> Link -> http://localhost:8000/openapi.json")
bullet(doc, "Postman genera automaticamente los 7 requests con sus schemas tipados")

doc.add_paragraph()
heading(doc, "13.4 Autenticacion (si API_KEY esta configurada)", 2, "1565C0")
bullet(doc, "Agregar header X-API-Key: <tu_api_key> en cada request a /predict, /analyze-content, /cache/stats o DELETE /cache")
bullet(doc, "En Swagger UI: boton 'Authorize' (candado) en la parte superior derecha")

# ── 14. CONCLUSION ────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "14. Conclusion", 1, "1E3A5F")
p = doc.add_paragraph(
    "Los 10 puntos del producto esperado estan implementados y verificados. "
    "La API es utilizable end-to-end desde Swagger UI, curl o Postman sin "
    "necesidad de leer el codigo fuente: contratos tipados y visibles en "
    "/docs, validacion de entrada explicita, manejo de errores en 3 capas "
    "(sub-servicio, HTTP, handler global), y 268 tests automatizados que "
    "cubren cada endpoint y cada codigo de error documentado."
)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Pendiente fuera de este alcance: ")
run.bold = True
run.font.size = Pt(9.5)
run2 = p.add_run(
    "commitear los cambios de esta sesion (README.md, backend/app/main.py, "
    "backend/app/api/routes.py, backend/app/schemas/, tests/, docs/api.md "
    "estan en el arbol de trabajo sin commitear)."
)
run2.font.size = Pt(9.5)

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Informe Final: API de IA Testeable desde un Cliente Externo  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "informe_final_api_ia.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
