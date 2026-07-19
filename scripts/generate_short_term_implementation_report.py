"""
Genera el informe de implementacion del plan de mejora - corto plazo
del AI Phishing Detector como .docx
Salida: docs/implementacion_plan_mejora_corto_plazo.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Implementacion del Plan de Mejora — Corto Plazo")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Seguimiento del diagnostico de arquitectura: riesgos cerrados")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. CONTEXTO ──────────────────────────────────────────────────────────
heading(doc, "1. Contexto", 1, "1E3A5F")
p = doc.add_paragraph(
    "El informe \"Diagnostico de Arquitectura: Componentes, Riesgos y Plan de Mejora\" "
    "identifico 4 acciones de corto plazo (1-2 semanas, bajo esfuerzo) para cerrar los "
    "riesgos de seguridad y operacion mas criticos con el menor costo de implementacion. "
    "Este documento reporta lo efectivamente implementado, verificado contra el codigo "
    "y la suite de tests."
)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_table(doc,
    ["Metrica", "Resultado"],
    [
        ["Tests ejecutados", "257 / 257 aprobados"],
        ["Riesgos de corto plazo cerrados", "4 / 4"],
        ["Archivos modificados", "6"],
        ["Regresiones detectadas", "0"],
    ],
    col_widths=[6.0, 8.5],
)

# ── 2. CAMBIOS IMPLEMENTADOS ─────────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Cambios Implementados", 1, "1E3A5F")

heading(doc, "2.1 Autenticacion obligatoria en produccion", 2, "2E7D32")
p = doc.add_paragraph(
    "Riesgo cerrado: el backend arrancaba sin autenticacion si API_KEY quedaba vacio, "
    "sin ninguna advertencia."
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "Nueva variable de entorno ENVIRONMENT (development | production)")
bullet(doc, "Si ENVIRONMENT=production y API_KEY esta vacio, el backend lanza RuntimeError y rehusa arrancar (fail-fast)")
bullet(doc, "Verificado manualmente: con ENVIRONMENT=production y API_KEY='' el import de backend.app.main levanta el error esperado; sin ENVIRONMENT (modo dev por defecto) /health responde 200 con normalidad")
add_code_block(doc,
"""backend/app/core/config.py
    environment: str = os.getenv("ENVIRONMENT", "development")

backend/app/main.py
    if settings.environment == "production" and not settings.api_key:
        raise RuntimeError("ENVIRONMENT=production requiere API_KEY configurada...")""")

doc.add_paragraph()
heading(doc, "2.2 Healthcheck en Docker Compose", 2, "2E7D32")
p = doc.add_paragraph(
    "Riesgo cerrado: un contenedor colgado (modelo corrupto, deadlock) no se detectaba "
    "automaticamente."
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "HEALTHCHECK sobre GET /health cada 30s, timeout 5s, 3 reintentos, start_period 15s")
bullet(doc, "Usa python -c con urllib (no requiere instalar curl en la imagen)")
add_code_block(doc,
"""docker-compose.yml
    healthcheck:
      test: ["CMD", "python", "-c",
             "import urllib.request as u; u.urlopen('http://localhost:8000/health', timeout=3)"]
      interval: 30s
      timeout: 5s
      retries: 3
      start_period: 15s""")

doc.add_paragraph()
heading(doc, "2.3 Logging de requests", 2, "2E7D32")
p = doc.add_paragraph(
    "Riesgo cerrado: no existia registro de auditoria minima (quien llamo, que endpoint, "
    "cuando, con que resultado)."
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "RequestLoggingMiddleware nuevo: registra metodo, path, status code, latencia en ms e IP del cliente")
bullet(doc, "Registrado como middleware mas externo -> tambien loguea las solicitudes bloqueadas por rate limiting (429)")
add_code_block(doc,
"""backend/app/main.py
    logger.info("%s %s %s %sms ip=%s",
                request.method, request.url.path, response.status_code, elapsed_ms, ip)

    # Orden: CORS -> RateLimitMiddleware -> RequestLoggingMiddleware (mas externo)""")

doc.add_paragraph()
heading(doc, "2.4 Documentacion del requisito de proxy TLS", 2, "2E7D32")
p = doc.add_paragraph(
    "Riesgo cerrado: no habia advertencia de que el backend expone HTTP plano y no debe "
    "exponerse directamente a internet."
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "Nueva seccion 'Despliegue en produccion' en README.md")
bullet(doc, "Indica: proxy TLS obligatorio (nginx / Caddy / Traefik), configurar ENVIRONMENT + API_KEY, revisar ALLOWED_ORIGINS")

doc.add_paragraph()
heading(doc, "2.5 Correccion adicional", 2, "1565C0")
p = doc.add_paragraph(
    "El README referenciaba backend/app/analyzers/test_html_analyzer.py, eliminado en la "
    "limpieza previa de codigo muerto. Se corrigio el comando de pruebas a pytest sin ruta "
    "especifica (ejecuta toda la suite via pytest.ini)."
)
p.runs[0].font.size = Pt(9.5)

# ── 3. ARCHIVOS MODIFICADOS ──────────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Archivos Modificados", 1, "1E3A5F")
add_table(doc,
    ["Archivo", "Cambio"],
    [
        ["backend/app/core/config.py", "Nueva variable settings.environment"],
        ["backend/app/main.py", "Guard de arranque en produccion + RequestLoggingMiddleware"],
        ["docker-compose.yml", "Bloque healthcheck sobre GET /health"],
        [".env.example", "Documentada la variable ENVIRONMENT"],
        ["README.md", "Seccion 'Despliegue en produccion' + fix de comando de pruebas"],
    ],
    col_widths=[5.5, 9.0],
)

# ── 4. VERIFICACION ──────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "4. Verificacion", 1, "1E3A5F")
add_table(doc,
    ["Prueba", "Resultado"],
    [
        ["Suite completa de pytest", "257 passed, 0 failed"],
        ["ENVIRONMENT=production + API_KEY vacio -> import backend.app.main", "RuntimeError lanzado como se esperaba"],
        ["Modo desarrollo (sin ENVIRONMENT) -> GET /health", "200 {'status': 'healthy'}"],
    ],
    col_widths=[8.0, 6.5],
)

doc.add_paragraph()
heading(doc, "5. Pendiente (mediano y largo plazo)", 2, "B71C1C")
p = doc.add_paragraph(
    "Estas acciones no se ejecutaron en este ciclo; siguen documentadas en el informe de "
    "diagnostico de arquitectura para una proxima iteracion:"
)
p.runs[0].font.size = Pt(9.5)
bullet(doc, "Mediano plazo: cache/rate-limit en Redis, multiples workers Uvicorn, CI con lint+tests, circuit breaker para APIs externas, decision sobre phishing_detector/")
bullet(doc, "Largo plazo: monitoreo de drift, pipeline de reentrenamiento, metricas de latencia por sub-servicio, consolidacion de tests duplicados")

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Implementacion del Plan de Mejora (Corto Plazo)  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "implementacion_plan_mejora_corto_plazo.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
