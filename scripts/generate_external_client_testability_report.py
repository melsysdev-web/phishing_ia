"""
Genera el informe "API testeable desde un cliente externo (Swagger / curl /
Postman)" del AI Phishing Detector como .docx
Salida: docs/informe_testabilidad_cliente_externo.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="C0C0C0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None,
              header_bg="1E3A5F", header_fg="FFFFFF",
              alt_bg="EEF2F7"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(header_fg)
        run.font.size = Pt(9)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            set_cell_bg(cell, bg)
            set_cell_border(cell, "D0D8E4")
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ══════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Phishing Detector")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run("Testeabilidad desde un Cliente Externo")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x54, 0x6E, 0x7A)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run("Swagger  ·  curl  ·  Postman  ·  Producto esperado")
run3.font.size = Pt(11)
run3.italic = True
run3.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = t4.add_run("Melvin Mendoza  ·  Rama phishing-ia/risk-engine-v2  ·  2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

doc.add_page_break()

# ── 1. REQUISITO ─────────────────────────────────────────────────────────
heading(doc, "1. Requisito", 1, "1E3A5F")
p = doc.add_paragraph(
    "“La API debe permitir que la funcionalidad de IA pueda ser probada desde "
    "un cliente externo como Swagger, curl, Postman o una interfaz web simple.”"
)
p.runs[0].italic = True
p.runs[0].font.size = Pt(10.5)

doc.add_paragraph()
add_table(doc,
    ["Cliente", "Estado", "Evidencia"],
    [
        ["Swagger UI", "Cumple",
         "FastAPI expone /docs automaticamente; con response_model + tags + summary/description en los 6 endpoints (ver informe_api_endpoints.docx) queda usable sin leer el codigo"],
        ["curl",       "Cumple",
         "Contrato JSON simple sobre HTTP; requiere header X-API-Key solo si API_KEY esta configurada"],
        ["Postman",    "Cumple",
         "Mismo contrato REST/JSON que curl; el spec OpenAPI (/openapi.json) se puede importar directo a Postman"],
        ["Interfaz web simple", "Cumple (via Swagger)",
         "No existe una UI standalone fuera de la extension de Chrome; Swagger UI cubre el caso de 'interfaz web' para pruebas manuales"],
    ],
    col_widths=[3.0, 2.5, 8.5],
)

# ── 2. PRODUCTO ESPERADO ─────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "2. Producto Esperado", 1, "1E3A5F")
p = doc.add_paragraph(
    "Estado de cada entregable, verificado contra el codigo y la documentacion "
    "actuales del repositorio (rama phishing-ia/risk-engine-v2)."
)
p.runs[0].font.size = Pt(9.5)
doc.add_paragraph()
add_table(doc,
    ["Entregable", "Estado", "Detalle"],
    [
        ["Codigo actualizado del proyecto", "Cumple",
         "Ultimo commit en la rama incluye Docker, seguridad, response_model/tags y la suite de tests ampliada"],
        ["Endpoint principal de IA", "Cumple",
         "POST /predict (pipeline completo) y POST /analyze-content (clasificacion de texto)"],
        ["Endpoint /health", "Cumple",
         "GET /health -> {\"status\": \"healthy\"}, sin autenticacion, usado por la extension y el HEALTHCHECK de Docker"],
        ["Endpoint /metadata o equivalente", "Cumple",
         "GET /metadata (nuevo): version de la API, disponibilidad de cada modelo ML, rate limit y config de cache — sin autenticacion, tag Sistema"],
        ["Contrato de entrada y salida", "Cumple",
         "Request: UrlRequest / TextRequest (Pydantic). Response: PredictResponse / ContentAnalysisResponse / CacheStatsResponse / CacheClearResponse / HealthResponse / RootResponse / MetadataResponse / ErrorResponse (backend/app/schemas/)"],
        ["Validacion basica de datos", "Cumple",
         "Pydantic valida tipos y el prefijo http(s):// de UrlRequest.url; errores devuelven 422 automatico con detalle de campo"],
        ["Manejo controlado de errores", "Cumple",
         "Exception handler global (@app.exception_handler(Exception) en main.py) agregado: cualquier excepcion no capturada en el flujo principal ahora devuelve 500 con contrato JSON consistente ({\"error\": ..., \"detail\": ...}), ademas de la degradacion por sub-servicio ya existente"],
        ["Evidencia de prueba", "Cumple",
         "268 tests (pytest.ini: testpaths = tests backend/tests), incluye tests de contrato para los 7 endpoints, 403, 429 y el 500 controlado"],
        ["README actualizado", "Cumple",
         "Nueva seccion 'Probar la API sin escribir codigo' (Swagger, Postman via /openapi.json, ejemplo curl); tabla de endpoints actualizada con /metadata; link a docs/api.md"],
        ["docs/api.md o equivalente", "Cumple",
         "Nuevo docs/api.md: contrato de los 7 endpoints, ejemplos curl, tabla de codigos de error (422/403/429/500), instrucciones de Swagger/Postman"],
    ],
    col_widths=[3.6, 2.0, 8.4],
)

# ── 3. DETALLE DE LO IMPLEMENTADO ─────────────────────────────────────────
doc.add_page_break()
heading(doc, "3. Detalle de lo Implementado", 1, "1E3A5F")

heading(doc, "3.1 Endpoint /metadata", 2, "2E7D32")
p = doc.add_paragraph(
    "GET /metadata expone informacion de la API en si misma, sin cargar ni "
    "inicializar los modelos (solo verifica existencia de archivos/carpetas "
    "en get_models_dir(), evitando romper el lazy loading de RoBERTa)."
)
p.runs[0].font.size = Pt(9.5)
add_code_block(doc,
"""GET /metadata
{
  "api_version": "1.0.0",
  "models": { "random_forest": true, "roberta_url": true, "roberta_content": true },
  "rate_limit_per_minute": 30,
  "cache_ttl_seconds": 600,
  "cache_max_size": 500
}""")

doc.add_paragraph()
heading(doc, "3.2 Manejo controlado de errores", 2, "2E7D32")
p = doc.add_paragraph("Estado por capa despues del cambio:")
p.runs[0].font.size = Pt(9.5)
bullet(doc, "APIs externas y sub-tareas del pipeline: sin cambios, ya degradaban a {\"error\": \"...\"} dentro de su propia clave (_safe(), try/except por servicio)")
bullet(doc, "Nuevo: @app.exception_handler(Exception) en main.py — cualquier excepcion que escape de esas rutas ahora devuelve 500 { \"error\": \"Error interno del servidor\", \"detail\": \"...\" }, con detail=null si ENVIRONMENT=production (no filtra detalles internos en despliegues publicos)")
bullet(doc, "Nota tecnica: Starlette re-lanza las excepciones de servidor en TestClient por diseno (para visibilidad en debugging/Sentry) aun cuando el handler ya envio la respuesta 500 correcta; el test usa TestClient(app, raise_server_exceptions=False) para verificar el contrato JSON real que recibe un cliente externo")

doc.add_paragraph()
heading(doc, "3.3 docs/api.md", 2, "2E7D32")
bullet(doc, "Nuevo documento Markdown versionado en git: referencia de los 7 endpoints, ejemplos curl, tabla de codigos de error, instrucciones para Swagger y Postman")
bullet(doc, "informe_api_endpoints.docx sigue existiendo como el informe de sesion; docs/api.md es la referencia estable que vive junto al codigo")

doc.add_paragraph()
heading(doc, "3.4 README — probar la API", 2, "2E7D32")
bullet(doc, "Nueva seccion 'Probar la API sin escribir codigo': link a /docs, instrucciones de import a Postman via /openapi.json, ejemplo curl completo")
bullet(doc, "Tabla de endpoints actualizada con /health y /metadata; nueva fila en 'Documentacion' apuntando a docs/api.md")

# ── 4. EVIDENCIA ──────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "4. Evidencia", 1, "1E3A5F")

heading(doc, "4.1 Ejemplo curl — POST /predict", 2, "2E7D32")
add_code_block(doc,
"""curl -X POST http://localhost:8000/predict \\
  -H "Content-Type: application/json" \\
  -H "X-API-Key: <tu_api_key>" \\
  -d '{"url": "https://ejemplo.com"}'""")

doc.add_paragraph()
heading(doc, "4.2 Ejemplo curl — POST /analyze-content", 2, "2E7D32")
add_code_block(doc,
"""curl -X POST http://localhost:8000/analyze-content \\
  -H "Content-Type: application/json" \\
  -H "X-API-Key: <tu_api_key>" \\
  -d '{"text": "..."}'   # >= 300 caracteres para ser clasificado""")

doc.add_paragraph()
heading(doc, "4.3 GET /metadata", 2, "2E7D32")
add_code_block(doc, 'curl http://localhost:8000/metadata')

doc.add_paragraph()
heading(doc, "4.4 Swagger UI y Postman", 2, "2E7D32")
bullet(doc, "http://localhost:8000/docs — interfaz interactiva generada por FastAPI, agrupada en tags Sistema / Analisis / Cache, con 'Try it out' funcional para los 7 endpoints")
bullet(doc, "http://localhost:8000/openapi.json — spec OpenAPI 3.1 importable directamente a Postman (Import > Link)")

doc.add_paragraph()
heading(doc, "4.5 Validacion de datos (ejemplo)", 2, "2E7D32")
add_code_block(doc,
"""POST /predict  { "url": "no-es-una-url" }
-> 422 { "detail": [{"loc": ["body","url"], "msg": "Value error, La URL debe
   comenzar con http:// o https://", "type": "value_error"}] }""")

doc.add_paragraph()
heading(doc, "4.6 Error 500 controlado (ejemplo)", 2, "2E7D32")
add_code_block(doc,
"""POST /predict  (excepcion no prevista en el flujo principal)
-> 500 { "error": "Error interno del servidor", "detail": "boom" }
   (detail=null si ENVIRONMENT=production)""")

# ── 5. RESULTADO ──────────────────────────────────────────────────────────
doc.add_page_break()
heading(doc, "5. Resultado", 1, "1E3A5F")
p = doc.add_paragraph(
    "Las 4 acciones de la version anterior de este informe (endpoint /metadata, "
    "exception handler global, docs/api.md, seccion de README) fueron implementadas "
    "y verificadas: 268/268 tests aprobados (265 previos + 3 nuevos: 2 para "
    "/metadata, 1 para el 500 controlado). El producto esperado del punto 2 "
    "queda completo."
)
p.runs[0].font.size = Pt(9.5)

# ── Pie de pagina ─────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("AI Phishing Detector — Testeabilidad desde un Cliente Externo  ·  Melvin Mendoza  ·  2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xA4, 0xAE)

# ── Guardar ─────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "informe_testabilidad_cliente_externo.docx")

doc.save(out_path)
print(f"Word guardado en: {os.path.abspath(out_path)}")
